import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

# --- DATA STRUCTURES ---

ITEMS_DATA = [
    {
        "id": 1,
        "name": "Blazer blanco doble botonadura",
        "desc_op1": "Blazer de tweed bouclé estructurado con doble botonadura, botones dorados y bolsillos funcionales. Colores: blanco y negro.",
        "colors_op1": "Blanco, Negro",
        "sizes_op1": "S, M, L",
        "price_op1": 40.80,
        "desc_op2": "Blazer de tweed bouclé estructurado con doble botonadura, botones dorados y bolsillos funcionales. Colores: negro y blanco.",
        "colors_op2": "Negro, Blanco",
        "sizes_op2": "S, M, L",
        "price_op2": 16.89,
        "img_op1": "1.jpeg",
        "img_op2": "1.jpeg",
        "potencial": "Excelente",
        "viabilidad": "Alta",
        "riesgo": "Bajo"
    },
    {
        "id": 2,
        "name": "Blazer rosado tweed",
        "desc_op1": "Blazer de tweed bouclé con corte princesa y doble botonadura dorada. Colores: rosa, blanco, negro, marrón, rojo y azul.",
        "colors_op1": "Rosa, Blanco, Negro, Marrón, Rojo, Azul",
        "sizes_op1": "S, M, L",
        "price_op1": 15.84,
        "desc_op2": "Blazer de tweed con corte princesa y doble botonadura dorada. Colores: rosa, blanco, negro, rojo y azul.",
        "colors_op2": "Rosa, Blanco, Negro, Rojo, Azul",
        "sizes_op2": "S, M, L",
        "price_op2": 11.59,
        "img_op1": "2.jpeg",
        "img_op2": "2.jpeg",
        "potencial": "Muy Alto",
        "viabilidad": "Alta",
        "riesgo": "Bajo"
    },
    {
        "id": 3,
        "name": "Blazer tweed cropped",
        "desc_op1": "Blazer de tweed estilo cropped con botones dorados. Colores: amarillo, azul, negro y blanco roto.",
        "colors_op1": "Amarillo, Azul, Negro, Blanco roto",
        "sizes_op1": "S, M, L",
        "price_op1": 18.00,
        "desc_op2": "Blazer de tweed estilo cropped con doble botonadura y botones dorados. Color: azul.",
        "colors_op2": "Azul",
        "sizes_op2": "S, M, L",
        "price_op2": 11.00,
        "img_op1": "3.jpeg",
        "img_op2": "3.jpeg",
        "potencial": "Medio-Alto",
        "viabilidad": "Alta",
        "riesgo": "Bajo"
    },
    {
        "id": 4,
        "name": "Chaqueta Chanel negra",
        "desc_op1": "Chaqueta corta de tweed estilo Chanel con botones metálicos. Color: negro.",
        "colors_op1": "Negro",
        "sizes_op1": "S, M, L",
        "price_op1": 18.00,
        "desc_op2": "Chaqueta corta de tweed estilo Chanel con botones dorados. Color: negro.",
        "colors_op2": "Negro",
        "sizes_op2": "S, M, L",
        "price_op2": 5.79,
        "img_op1": "4.jpeg",
        "img_op2": "4.jpeg",
        "potencial": "Muy Alto",
        "viabilidad": "Excelente",
        "riesgo": "Muy Bajo"
    },
    {
        "id": 5,
        "name": "Chaqueta tweed francés",
        "desc_op1": "Chaqueta corta de tweed francés con botones de perla y forro satinado. Color: verde.",
        "colors_op1": "Verde",
        "sizes_op1": "S, M, L",
        "price_op1": 21.84,
        "desc_op2": "Chaqueta corta de tweed francés con botones de perla y forro satinado. Color: azul.",
        "colors_op2": "Azul",
        "sizes_op2": "S, M, L",
        "price_op2": 11.78,
        "img_op1": "5.jpeg",
        "img_op2": "5.jpeg",
        "potencial": "Alto",
        "viabilidad": "Alta",
        "riesgo": "Bajo"
    },
    {
        "id": 6,
        "name": "Abrigo camel largo",
        "desc_op1": "Abrigo largo con doble botonadura dorada y corte princesa. Colores: caqui, gris, negro, naranja, verde militar, burdeos, azul y celeste.",
        "colors_op1": "Caqui, Gris, Negro, Naranja, Burdeos, Azul, etc.",
        "sizes_op1": "S, M, L, XL",
        "price_op1": 13.20,
        "desc_op2": "Abrigo largo con doble botonadura dorada y corte princesa. Colores: blanco, caqui, negro y gris.",
        "colors_op2": "Blanco, Caqui, Negro, Gris",
        "sizes_op2": "S, M, L, XL",
        "price_op2": 8.25,
        "img_op1": "6.jpeg",
        "img_op2": "6.jpeg",
        "potencial": "Alto",
        "viabilidad": "Excelente",
        "riesgo": "Bajo"
    },
    {
        "id": 7,
        "name": "Chaqueta blazer ribetes",
        "desc_op1": "Chaqueta blazer estructurada con ribetes en contraste y cinturón decorativo. Colores: blanco y negro.",
        "colors_op1": "Blanco, Negro",
        "sizes_op1": "S, M, L",
        "price_op1": 21.84,
        "desc_op2": "Chaqueta blazer estructurada con ribetes en contraste y cinturón decorativo. Colores: negro y blanco.",
        "colors_op2": "Negro, Blanco",
        "sizes_op2": "S, M, L",
        "price_op2": 15.51,
        "img_op1": "7.jpeg",
        "img_op2": "7.jpeg",
        "potencial": "Medio",
        "viabilidad": "Media",
        "riesgo": "Bajo"
    },
    {
        "id": 8,
        "name": "Blazer doble botonadura moderno",
        "desc_op1": "Chaqueta blazer con cinturón ajustable y doble botonadura de estilo moderno. Diseño estampado.",
        "colors_op1": "Diseño Estampado",
        "sizes_op1": "S, M, L",
        "price_op1": 31.20,
        "desc_op2": "Chaqueta blazer con cinturón ajustable y doble botonadura de estilo moderno. Color: negro.",
        "colors_op2": "Negro",
        "sizes_op2": "S, M, L",
        "price_op2": 6.09,
        "img_op1": "8.jpeg",
        "img_op2": "8opcion 2.jpg",
        "potencial": "Muy Alto",
        "viabilidad": "Excelente",
        "riesgo": "Muy Bajo"
    },
    {
        "id": 9,
        "name": "Cárdigan perlas decorativas",
        "desc_op1": "Cárdigan corto de punto con perlas decorativas. Colores: rojo, rosa, flor de cerezo, amarillo, beige y azul.",
        "colors_op1": "Rojo, Rosa, Cerezo, Amarillo, Beige, Azul",
        "sizes_op1": "S, M, L",
        "price_op1": 18.00,
        "desc_op2": "Cárdigan corto de punto con perlas decorativas. Colores: blanco, rosa, amarillo, verde, azul y negro.",
        "colors_op2": "Blanco, Rosa, Amarillo, Verde, Azul, Negro",
        "sizes_op2": "Talla Única",
        "price_op2": 7.85,
        "img_op1": "9.jpeg",
        "img_op2": "9.jpeg",
        "potencial": "Alto",
        "viabilidad": "Alta",
        "riesgo": "Bajo"
    },
    {
        "id": 10,
        "name": "Cárdigan punto texturizado",
        "desc_op1": "Cárdigan de punto texturizado con ribetes negros y botones decorativos. Color: blanco.",
        "colors_op1": "Blanco",
        "sizes_op1": "S, M, L",
        "price_op1": 9.60,
        "desc_op2": "Cárdigan de punto texturizado con ribetes negros y botones decorativos. Colores: negro y blanco.",
        "colors_op2": "Negro, Blanco",
        "sizes_op2": "S, M, L",
        "price_op2": 7.07,
        "img_op1": "10.jpeg",
        "img_op2": "10.jpeg",
        "potencial": "Muy Alto",
        "viabilidad": "Excelente",
        "riesgo": "Muy Bajo"
    },
    {
        "id": 11,
        "name": "Cárdigan punto acanalado",
        "desc_op1": "Cárdigan tejido de punto acanalado con cuello marinero y botones decorativos. Colores: blanco y rojo.",
        "colors_op1": "Blanco, Rojo",
        "sizes_op1": "S, M, L",
        "price_op1": 11.70,
        "desc_op2": "Cárdigan tejido de punto acanalado con botones de perla. Colores: rosa, albaricoque, amarillo, rojo y negro.",
        "colors_op2": "Rosa, Albaricoque, Amarillo, Rojo, Negro",
        "sizes_op2": "S, M, L",
        "price_op2": 4.52,
        "img_op1": "11.jpeg",
        "img_op2": "11.jpeg",
        "potencial": "Muy Alto",
        "viabilidad": "Excelente",
        "riesgo": "Bajo"
    },
    {
        "id": 12,
        "name": "Cárdigan con volantes",
        "desc_op1": "Cárdigan con volantes y botones metálicos decorativos. Colores: marrón, café, azul cielo, rosa, blanco, negro y rojo.",
        "colors_op1": "Marrón, Café, Celeste, Rosa, Blanco, Negro, Rojo",
        "sizes_op1": "S, M, L",
        "price_op1": 12.00,
        "desc_op2": "Cárdigan con volantes y botones metálicos decorativos. Colores: rosa, negro, blanco, rojo y azul cielo.",
        "colors_op2": "Rosa, Negro, Blanco, Rojo, Celeste",
        "sizes_op2": "S, M, L",
        "price_op2": 9.43,
        "img_op1": "12.jpeg",
        "img_op2": "12.jpeg",
        "potencial": "Medio-Alto",
        "viabilidad": "Alta",
        "riesgo": "Bajo"
    },
    {
        "id": 13,
        "name": "Blusa entallada manga larga",
        "desc_op1": "Blusa entallada de manga larga con volantes y botones. Colores: blanco roto, rojo y azul oscuro.",
        "colors_op1": "Blanco roto, Rojo, Azul oscuro",
        "sizes_op1": "S, M, L",
        "price_op1": 10.80,
        "desc_op2": "Prenda de poliéster. Color según imagen.",
        "colors_op2": "Según imagen",
        "sizes_op2": "S, M, L",
        "price_op2": 3.20,
        "img_op1": "13.jpeg",
        "img_op2": "13opcion 2.jpg",
        "potencial": "Medio (Complementario)",
        "viabilidad": "Alta",
        "riesgo": "Bajo"
    },
    {
        "id": 14,
        "name": "Chaqueta biker Balmain",
        "desc_op1": "Chaqueta biker de cuero sintético estilo Balmain con botones dorados decorativos. Material: PU.",
        "colors_op1": "Marrón / Negro (según imagen)",
        "sizes_op1": "S, M, L",
        "price_op1": 32.40,
        "desc_op2": "Chaqueta biker de cuero sintético estilo Balmain con botones dorados. Color: negro.",
        "colors_op2": "Negro",
        "sizes_op2": "S, M, L",
        "price_op2": 52.63,
        "img_op1": "14.jpeg",
        "img_op2": "14.jpeg",
        "potencial": "Bajo-Medio (Alto Costo)",
        "viabilidad": "Baja",
        "riesgo": "Medio-Alto"
    },
    {
        "id": 15,
        "name": "Termo térmico 40 oz",
        "desc_op1": "Termo térmico de 40 oz con aislamiento al vacío y tapa hermética. Color: rosa.",
        "colors_op1": "Rosa",
        "sizes_op1": "40 oz",
        "price_op1": 14.40,
        "desc_op2": "Termo térmico de 40 oz con aislamiento al vacío y tapa hermética. Colores: azul y rosa.",
        "colors_op2": "Azul, Rosa",
        "sizes_op2": "40 oz",
        "price_op2": 8.44,
        "img_op1": "15.jpeg",
        "img_op2": "15.jpeg",
        "potencial": "Alto (Complementario)",
        "viabilidad": "Alta",
        "riesgo": "Bajo"
    }
]

# --- CORPORATE COLOR PALETTE ---
COLOR_MIDNIGHT_NAVY = RGBColor(0x0A, 0x16, 0x28)  # #0A1628
COLOR_COPPER = RGBColor(0xC1, 0x7F, 0x3E)         # #C17F3E
COLOR_GOLD_LIGHT = RGBColor(0xDB, 0xA0, 0x4E)     # #DBA04E
COLOR_WARM_CREAM = RGBColor(0xF5, 0xF0, 0xE8)     # #F5F0E8
COLOR_STEEL_GRAY = RGBColor(0x88, 0x92, 0xA0)     # #8892A0
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_TEXT_DARK = RGBColor(0x22, 0x28, 0x31)

HEX_MIDNIGHT_NAVY = "0A1628"
HEX_COPPER = "C17F3E"
HEX_GOLD_LIGHT = "DBA04E"
HEX_WARM_CREAM = "F5F0E8"
HEX_STEEL_GRAY = "8892A0"
HEX_WHITE = "FFFFFF"

# --- HELPER FUNCTIONS FOR STYLING ---

def set_font(run, font_name="Calibri", size_pt=11, bold=False, color_rgb=None, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def set_para_format(paragraph, before_pt=0, after_pt=6, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.LEFT, keep_with_next=False):
    paragraph.paragraph_format.space_before = Pt(before_pt)
    paragraph.paragraph_format.space_after = Pt(after_pt)
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.alignment = alignment
    paragraph.paragraph_format.keep_with_next = keep_with_next

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_cell_valign(cell, valign="center"):
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{valign}"/>')
    tcPr.append(vAlign)

def set_cell_borders(cell, top="none", left="none", bottom="none", right="none", color="CCCCCC", size="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    top_xml = f'<w:top w:val="{top}" w:sz="{size}" w:space="0" w:color="{color}"/>' if top != "none" else '<w:top w:val="none"/>'
    left_xml = f'<w:left w:val="{left}" w:sz="{size}" w:space="0" w:color="{color}"/>' if left != "none" else '<w:left w:val="none"/>'
    bottom_xml = f'<w:bottom w:val="{bottom}" w:sz="{size}" w:space="0" w:color="{color}"/>' if bottom != "none" else '<w:bottom w:val="none"/>'
    right_xml = f'<w:right w:val="{right}" w:sz="{size}" w:space="0" w:color="{color}"/>' if right != "none" else '<w:right w:val="none"/>'
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}>{top_xml}{left_xml}{bottom_xml}{right_xml}</w:tcBorders>')
    tcPr.append(borders)

def set_table_borders(table, color="D3D3D3", size="4"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/><w:left w:val="none"/><w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/><w:right w:val="none"/><w:insideH w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/><w:insideV w:val="none"/></w:tblBorders>')
    tblPr.append(borders)

def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>')
    tblPr.append(borders)

def add_heading_styled(doc, text, level=1, before_pt=18, after_pt=6):
    p = doc.add_paragraph()
    set_para_format(p, before_pt=before_pt, after_pt=after_pt, alignment=WD_ALIGN_PARAGRAPH.LEFT, keep_with_next=True)
    
    font_size = 18 if level == 1 else (14 if level == 2 else 12)
    font_bold = True
    font_color = COLOR_MIDNIGHT_NAVY if level == 1 else (COLOR_COPPER if level == 2 else COLOR_STEEL_GRAY)
    
    # Prefix number placeholder can be added manually or automatically.
    run = p.add_run(text)
    set_font(run, "Montserrat SemiBold", font_size, font_bold, font_color)
    
    # Set style name for TOC mapping
    p.style = doc.styles[f'Heading {level}']
    return p

def add_bullet_point(doc, bold_lead, text_body):
    p = doc.add_paragraph()
    set_para_format(p, before_pt=2, after_pt=4)
    p.paragraph_format.left_indent = Inches(0.25)
    
    r_bullet = p.add_run("•  ")
    set_font(r_bullet, "Calibri", 11, True, COLOR_COPPER)
    
    if bold_lead:
        r_lead = p.add_run(bold_lead + ": ")
        set_font(r_lead, "Montserrat Medium", 10.5, True, COLOR_MIDNIGHT_NAVY)
        
    r_body = p.add_run(text_body)
    set_font(r_body, "Calibri", 10.5, False, COLOR_TEXT_DARK)

def create_callout_box(doc, title, text, border_color=HEX_COPPER, bg_color=HEX_WARM_CREAM):
    table = doc.add_table(1, 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(7.0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    set_cell_borders(cell, left="single", color=border_color, size="36") # 4.5pt thick border
    
    p1 = cell.paragraphs[0]
    set_para_format(p1, before_pt=0, after_pt=4)
    r1 = p1.add_run(title)
    set_font(r1, "Montserrat SemiBold", 11, True, COLOR_COPPER)
    
    p2 = cell.add_paragraph()
    set_para_format(p2, before_pt=0, after_pt=0)
    r2 = p2.add_run(text)
    set_font(r2, "Calibri", 10.5, False, COLOR_MIDNIGHT_NAVY)
    
    # Add an empty paragraph after the table to restore cursor flow
    doc.add_paragraph()

def add_page_number_to_footer(run):
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_total_pages_to_footer(run):
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_table_of_contents(paragraph):
    run = paragraph.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve">TOC \\o "1-3" \\h \\z \\u</w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def enable_update_fields(doc):
    element = parse_xml(f'<w:updateFields {nsdecls("w")} w:val="true"/>')
    doc.settings.element.append(element)

# --- MAIN GENERATOR FUNCTION ---

def generate_docx():
    print("Initializing document...")
    doc = Document()
    enable_update_fields(doc)
    
    # Set default style to Calibri
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = COLOR_TEXT_DARK
    
    # ---------------------------------------------
    # SECTION 0: PORTADA (COVER PAGE)
    # ---------------------------------------------
    print("Creating Cover Page...")
    section_cover = doc.sections[0]
    section_cover.page_width = Inches(8.5)
    section_cover.page_height = Inches(11.0)
    # Zero margins for full page background bleed
    section_cover.top_margin = Inches(0)
    section_cover.bottom_margin = Inches(0)
    section_cover.left_margin = Inches(0)
    section_cover.right_margin = Inches(0)
    
    # 1x1 Table to cover the page
    cover_table = doc.add_table(1, 1)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_table.autofit = False
    
    row = cover_table.rows[0]
    row.height = Inches(11.0)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    cell = cover_table.cell(0, 0)
    cell.width = Inches(8.5)
    set_cell_background(cell, HEX_MIDNIGHT_NAVY)
    set_cell_margins(cell, top=720, bottom=500, left=720, right=720) # 0.5" margin inside cell
    set_cell_valign(cell, "center")
    remove_table_borders(cover_table)
    
    # Inside Cover Content
    p_logo = cell.paragraphs[0]
    set_para_format(p_logo, before_pt=0, after_pt=24, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    logo_path = "LOGO1 400x83 (1).png"
    if os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, width=Inches(3.2))
    else:
        r_logo_fail = p_logo.add_run("[ SMART IMPORT TRADE ]")
        set_font(r_logo_fail, "Montserrat SemiBold", 18, True, COLOR_GOLD_LIGHT)
        
    p_title = cell.add_paragraph()
    set_para_format(p_title, before_pt=18, after_pt=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    r_title = p_title.add_run("INFORME PRELIMINAR DE SOURCING Y EVALUACIÓN COMERCIAL")
    set_font(r_title, "Montserrat SemiBold", 22, True, COLOR_WHITE)
    
    p_sub = cell.add_paragraph()
    set_para_format(p_sub, before_pt=0, after_pt=24, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    r_sub = p_sub.add_run("Proyecto: Desarrollo de Colección de Blazers, Chaquetas Tweed y Cárdigans")
    set_font(r_sub, "Montserrat Medium", 12, False, COLOR_GOLD_LIGHT)
    
    # Visual Collage (Table 1 row, 3 columns)
    collage_table = cell.add_table(1, 3)
    collage_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(collage_table)
    
    col_images = ["1.jpeg", "2.jpeg", "6.jpeg"]
    for i, img_name in enumerate(col_images):
        col_cell = collage_table.cell(0, i)
        col_cell.width = Inches(2.2)
        set_cell_margins(col_cell, top=20, bottom=20, left=40, right=40)
        set_cell_background(col_cell, HEX_MIDNIGHT_NAVY)
        set_cell_valign(col_cell, "center")
        
        img_p = col_cell.paragraphs[0]
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_path = os.path.join("fotos blazer", img_name)
        if os.path.exists(img_path):
            img_p.add_run().add_picture(img_path, width=Inches(1.95))
        else:
            r_fail = img_p.add_run(f"[Foto {i+1}]")
            set_font(r_fail, "Calibri", 10, False, COLOR_STEEL_GRAY)
            
    # Metadata Box at the bottom of cover
    p_meta_space = cell.add_paragraph()
    set_para_format(p_meta_space, before_pt=24, after_pt=12)
    
    meta_table = cell.add_table(5, 2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(meta_table)
    
    meta_data = [
        ("Cliente:", "Importaciones Manorca S.A.C."),
        ("RUC:", "20615677010"),
        ("Atención:", "Gisela y Wilfredo"),
        ("Fecha:", "08 de junio de 2026"),
        ("Estatus:", "Documento Confidencial – Uso Exclusivo del Cliente")
    ]
    
    for idx, (label, val) in enumerate(meta_data):
        row_cell_lbl = meta_table.cell(idx, 0)
        row_cell_lbl.width = Inches(2.0)
        set_cell_background(row_cell_lbl, HEX_MIDNIGHT_NAVY)
        p_lbl = row_cell_lbl.paragraphs[0]
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_lbl = p_lbl.add_run(label + "  ")
        set_font(r_lbl, "Montserrat Medium", 9.5, True, COLOR_STEEL_GRAY)
        
        row_cell_val = meta_table.cell(idx, 1)
        row_cell_val.width = Inches(4.5)
        set_cell_background(row_cell_val, HEX_MIDNIGHT_NAVY)
        p_val = row_cell_val.paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_val = p_val.add_run(val)
        if idx == 4:
            set_font(r_val, "Calibri", 9.5, False, COLOR_COPPER, italic=True)
        else:
            set_font(r_val, "Calibri", 10, False, COLOR_WHITE)
            
    # ---------------------------------------------
    # SECTION 1: CARTA DE PRESENTACIÓN & TOC
    # ---------------------------------------------
    print("Creating Content Section and Presentation Letter...")
    section_content = doc.add_section(docx.enum.section.WD_SECTION.NEW_PAGE)
    section_content.header.is_linked_to_previous = False
    section_content.footer.is_linked_to_previous = False
    
    # 0.75" Margins for Content
    section_content.top_margin = Inches(0.75)
    section_content.bottom_margin = Inches(0.75)
    section_content.left_margin = Inches(0.75)
    section_content.right_margin = Inches(0.75)
    section_content.page_width = Inches(8.5)
    section_content.page_height = Inches(11.0)
    
    # Custom Header for Content
    header_table = section_content.header.add_table(1, 2, Inches(7.0))
    remove_table_borders(header_table)
    # Left Header
    cell_hl = header_table.cell(0, 0)
    p_hl = cell_hl.paragraphs[0]
    p_hl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_hl = p_hl.add_run("SMART IMPORT TRADE | Sourcing Internacional")
    set_font(r_hl, "Montserrat Medium", 8, False, COLOR_STEEL_GRAY)
    # Right Header
    cell_hr = header_table.cell(0, 1)
    p_hr = cell_hr.paragraphs[0]
    p_hr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_hr = p_hr.add_run("Colección Blazers, Tweed & Cárdigans 2026")
    set_font(r_hr, "Montserrat Medium", 8, False, COLOR_STEEL_GRAY)
    
    # Custom Footer for Content
    footer_table = section_content.footer.add_table(1, 2, Inches(7.0))
    remove_table_borders(footer_table)
    # Left Footer
    cell_fl = footer_table.cell(0, 0)
    p_fl = cell_fl.paragraphs[0]
    p_fl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_fl = p_fl.add_run("Confidencial | Importaciones Manorca S.A.C.")
    set_font(r_fl, "Calibri", 8, False, COLOR_STEEL_GRAY)
    # Right Footer
    cell_fr = footer_table.cell(0, 1)
    p_fr = cell_fr.paragraphs[0]
    p_fr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_fr = p_fr.add_run("Página ")
    set_font(r_fr, "Calibri", 8, False, COLOR_STEEL_GRAY)
    add_page_number_to_footer(p_fr.add_run())
    r_fr2 = p_fr.add_run(" de ")
    set_font(r_fr2, "Calibri", 8, False, COLOR_STEEL_GRAY)
    add_total_pages_to_footer(p_fr.add_run())
    
    # Presentation Letter Heading
    p_pres_title = doc.add_paragraph()
    set_para_format(p_pres_title, before_pt=12, after_pt=12)
    r_pres_title = p_pres_title.add_run("CARTA DE PRESENTACIÓN")
    set_font(r_pres_title, "Montserrat SemiBold", 16, True, COLOR_MIDNIGHT_NAVY)
    
    p_dest = doc.add_paragraph()
    set_para_format(p_dest, before_pt=6, after_pt=12)
    r_dest = p_dest.add_run("Estimados Gisela y Wilfredo (Importaciones Manorca S.A.C.):")
    set_font(r_dest, "Montserrat Medium", 11, True, COLOR_MIDNIGHT_NAVY)
    
    p_body1 = doc.add_paragraph()
    set_para_format(p_body1, before_pt=0, after_pt=6)
    r_body1 = p_body1.add_run("Nos complace presentar el presente ")
    set_font(r_body1, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    r_body1_bold = p_body1.add_run("Informe Preliminar de Sourcing y Evaluación Comercial")
    set_font(r_body1_bold, "Calibri", 10.5, True, COLOR_MIDNIGHT_NAVY)
    r_body1_end = p_body1.add_run(" desarrollado por Smart Import Trade. El objetivo de este documento es consolidar el trabajo de búsqueda, evaluación y comparación de proveedores realizado para la identificación de alternativas viables de importación de blazers, chaquetas tipo tweed y cárdigans para el mercado peruano.")
    set_font(r_body1_end, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    
    p_body2 = doc.add_paragraph()
    set_para_format(p_body2, before_pt=0, after_pt=6)
    r_body2 = p_body2.add_run("Durante este proceso se analizaron diferentes proveedores, modelos, estructuras de precios y posibilidades comerciales con la finalidad de identificar opciones que permitan optimizar la inversión inicial y reducir el riesgo comercial de una primera importación de prueba.")
    set_font(r_body2, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    
    p_body3 = doc.add_paragraph()
    set_para_format(p_body3, before_pt=0, after_pt=18)
    r_body3 = p_body3.add_run("El presente informe resume los hallazgos obtenidos, las alternativas identificadas y las recomendaciones preliminares para la siguiente etapa del proyecto.")
    set_font(r_body3, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    
    p_close = doc.add_paragraph()
    set_para_format(p_close, before_pt=12, after_pt=24)
    r_close = p_close.add_run("Atentamente,\n")
    set_font(r_close, "Calibri", 11, False, COLOR_TEXT_DARK)
    r_close_firm = p_close.add_run("SMART IMPORT TRADE")
    set_font(r_close_firm, "Montserrat SemiBold", 11, True, COLOR_COPPER)
    
    # TABLE OF CONTENTS
    p_toc_title = doc.add_paragraph()
    set_para_format(p_toc_title, before_pt=24, after_pt=6)
    r_toc_title = p_toc_title.add_run("ÍNDICE GENERAL")
    set_font(r_toc_title, "Montserrat Medium", 13, True, COLOR_COPPER)
    
    p_toc = doc.add_paragraph()
    set_para_format(p_toc, before_pt=6, after_pt=6)
    add_table_of_contents(p_toc)
    
    p_toc_note = doc.add_paragraph()
    set_para_format(p_toc_note, before_pt=0, after_pt=12)
    r_toc_note = p_toc_note.add_run("*(Para visualizar el índice, haga clic derecho sobre el campo anterior en Word y seleccione 'Actualizar campos' o presione F9).*")
    set_font(r_toc_note, "Calibri", 9, False, COLOR_STEEL_GRAY, italic=True)
    
    # ---------------------------------------------
    # SECTION 2: RESUMEN EJECUTIVO (Page 3)
    # ---------------------------------------------
    print("Creating Resumen Ejecutivo...")
    doc.add_page_break()
    
    add_heading_styled(doc, "1. RESUMEN EJECUTIVO", level=1)
    
    p_re1 = doc.add_paragraph()
    set_para_format(p_re1, before_pt=0, after_pt=6)
    r_re1 = p_re1.add_run("El presente informe consolida el avance preliminar del proceso de búsqueda, evaluación y comparación de proveedores para la importación de blazers, chaquetas tipo tweed y cárdigans solicitado por ")
    set_font(r_re1, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    r_re1_bold = p_re1.add_run("Importaciones Manorca S.A.C.")
    set_font(r_re1_bold, "Calibri", 10.5, True, COLOR_MIDNIGHT_NAVY)
    
    p_re2 = doc.add_paragraph()
    set_para_format(p_re2, before_pt=0, after_pt=6)
    r_re2 = p_re2.add_run("Como resultado del exhaustivo trabajo de sourcing realizado por nuestro equipo, se lograron preseleccionar ")
    set_font(r_re2, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    r_re2_bold = p_re2.add_run("dos proveedores internacionales")
    set_font(r_re2_bold, "Calibri", 10.5, True, COLOR_MIDNIGHT_NAVY)
    r_re2_end = p_re2.add_run(" con propuestas comerciales viables. Ambos proveedores ofrecen colecciones con excelente similitud visual a los modelos de referencia, pero presentan estructuras de costos, calidades, mínimos de compra (MOQ) y niveles de riesgo financiero marcadamente distintos.")
    set_font(r_re2_end, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    
    p_re3 = doc.add_paragraph()
    set_para_format(p_re3, before_pt=0, after_pt=12)
    r_re3 = p_re3.add_run("Se evaluaron un total de 15 referencias de productos, que abarcan blazers de tweed, abrigos y cárdigans, además de artículos complementarios como blusas y termos térmicos. Luego de analizar detalladamente las cotizaciones de ambos fabricantes, se concluye que existe una oportunidad excepcional de ahorro y optimización presupuestaria eligiendo la opción de menor costo unitario.")
    set_font(r_re3, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    
    create_callout_box(
        doc, 
        "CONCLUSIÓN PRELIMINAR SMART IMPORT TRADE", 
        "La Opción 2 (Proveedor ID-MU) representa actualmente la alternativa más competitiva para desarrollar una importación inicial piloto. Sus precios significativamente más bajos permiten construir una colección completa de prueba respetando el presupuesto estimado de USD 5,000, minimizando la exposición al riesgo comercial y maximizando los márgenes potenciales de venta en el mercado peruano."
    )
    
    # ---------------------------------------------
    # SECTION 3: METODOLOGÍA (Page 4)
    # ---------------------------------------------
    print("Creating Metodología...")
    doc.add_page_break()
    add_heading_styled(doc, "2. METODOLOGÍA DE BÚSQUEDA Y EVALUACIÓN", level=1)
    
    p_metod_intro = doc.add_paragraph()
    set_para_format(p_metod_intro, before_pt=0, after_pt=12)
    r_metod_intro = p_metod_intro.add_run("Para garantizar una propuesta comercial sólida, Smart Import Trade aplicó una metodología estructurada de sourcing y análisis comercial basada en seis fases estratégicas:")
    set_font(r_metod_intro, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    
    add_bullet_point(doc, "Investigación y búsqueda de proveedores", "Filtro inicial en plataformas de comercio exterior de fabricantes especializados en tejido de punto, tweed y sastrería de alta gama con certificaciones de exportación.")
    add_bullet_point(doc, "Comparación de precios y MOQ", "Cálculo detallado de costos unitarios y requerimientos mínimos de compra (Minimum Order Quantity - MOQ) para evaluar el impacto en el flujo de caja del cliente.")
    add_bullet_point(doc, "Evaluación visual de similitud", "Validación técnica de que los modelos ofertados coincidan en corte, textura y botones con las referencias de diseño solicitadas por Importaciones Manorca S.A.C.")
    add_bullet_point(doc, "Análisis de potencial comercial", "Evaluación de la rotación esperada de cada artículo en el mercado peruano basándose en tendencias locales de moda y versatilidad de las prendas.")
    add_bullet_point(doc, "Evaluación de riesgo financiero", "Comparación de la inversión total requerida por cada opción frente al presupuesto disponible, priorizando la diversificación del portafolio.")
    add_bullet_point(doc, "Identificación de oportunidades de optimización", "Búsqueda de artículos complementarios con alta demanda que compartan el espacio de importación para diluir costos logísticos fijos.")

    p_metod_concl = doc.add_paragraph()
    set_para_format(p_metod_concl, before_pt=12, after_pt=6)
    r_metod_concl = p_metod_concl.add_run("A continuación se detallan las propuestas recibidas de cada uno de los proveedores preseleccionados.")
    set_font(r_metod_concl, "Calibri", 10.5, False, COLOR_TEXT_DARK)

    # ---------------------------------------------
    # SECTION 4: OPCIÓN 1 – PROVEEDOR ID-EG (Page 5)
    # ---------------------------------------------
    print("Creating Opción 1...")
    doc.add_page_break()
    add_heading_styled(doc, "3. ALTERNATIVA IDENTIFICADA: OPCIÓN 1 – PROVEEDOR ID-EG", level=1)
    
    p_op1_desc = doc.add_paragraph()
    set_para_format(p_op1_desc, before_pt=0, after_pt=6)
    r_op1_desc = p_op1_desc.add_run("Esta alternativa presenta productos orientados a un segmento de precio medio-alto. Destaca por acabados premium y una amplia gama de opciones cromáticas, aunque exige un desembolso significativamente superior.")
    set_font(r_op1_desc, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    
    # Small stats panel
    table_op1_stats = doc.add_table(1, 3)
    table_op1_stats.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table_op1_stats)
    stats_data_op1 = [
        ("INVERSIÓN TOTAL", "$57,924.00", HEX_MIDNIGHT_NAVY),
        ("CANTIDAD TOTAL", "3,000 unidades", HEX_STEEL_GRAY),
        ("PROMEDIO PRECIO", "$19.31 / unidad", HEX_COPPER)
    ]
    for idx, (title, value, color) in enumerate(stats_data_op1):
        cell_stat = table_op1_stats.cell(0, idx)
        cell_stat.width = Inches(2.3)
        set_cell_background(cell_stat, HEX_WARM_CREAM)
        set_cell_margins(cell_stat, top=80, bottom=80, left=100, right=100)
        set_cell_borders(cell_stat, top="single", color=HEX_STEEL_GRAY, size="4")
        p_stat_t = cell_stat.paragraphs[0]
        p_stat_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_stat_t = p_stat_t.add_run(title + "\n")
        set_font(r_stat_t, "Montserrat Medium", 8.5, False, COLOR_STEEL_GRAY)
        r_stat_v = p_stat_t.add_run(value)
        set_font(r_stat_v, "Montserrat SemiBold", 12, True, RGBColor(int(color[0:2],16), int(color[2:4],16), int(color[4:6],16)))
        
    doc.add_paragraph() # spacer
    
    # 15 Product Table for Op 1
    table_op1 = doc.add_table(16, 6)
    table_op1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_op1, color="CCCCCC", size="4")
    
    headers_op1 = ["Item", "Foto", "Descripción / Especificaciones", "Colores / Tallas", "Precio", "Total (200 u.)"]
    widths_op1 = [Inches(0.4), Inches(0.95), Inches(2.45), Inches(1.4), Inches(0.8), Inches(1.0)]
    
    # Header Row Style
    hdr_row = table_op1.rows[0]
    hdr_row.height = Inches(0.4)
    hdr_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for i, title in enumerate(headers_op1):
        cell_hdr = hdr_row.cells[i]
        cell_hdr.width = widths_op1[i]
        set_cell_background(cell_hdr, HEX_MIDNIGHT_NAVY)
        set_cell_margins(cell_hdr, top=100, bottom=100, left=100, right=100)
        set_cell_valign(cell_hdr, "center")
        p_hdr = cell_hdr.paragraphs[0]
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0,1,4,5] else WD_ALIGN_PARAGRAPH.LEFT
        r_hdr = p_hdr.add_run(title)
        set_font(r_hdr, "Montserrat SemiBold", 9, True, COLOR_WHITE)
        
    # Data Rows
    for idx, item in enumerate(ITEMS_DATA):
        row_idx = idx + 1
        data_row = table_op1.rows[row_idx]
        data_row.height = Inches(0.95)
        data_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        
        bg_row_color = HEX_WHITE if idx % 2 == 0 else HEX_WARM_CREAM
        
        # Col 0: Item
        c0 = data_row.cells[0]
        c0.width = widths_op1[0]
        set_cell_background(c0, bg_row_color)
        set_cell_valign(c0, "center")
        p_c0 = c0.paragraphs[0]
        p_c0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c0 = p_c0.add_run(str(item["id"]))
        set_font(r_c0, "Montserrat SemiBold", 10, True, COLOR_MIDNIGHT_NAVY)
        
        # Col 1: Foto
        c1 = data_row.cells[1]
        c1.width = widths_op1[1]
        set_cell_background(c1, bg_row_color)
        set_cell_valign(c1, "center")
        p_c1 = c1.paragraphs[0]
        p_c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_path = os.path.join("fotos blazer", item["img_op1"])
        if os.path.exists(img_path):
            p_c1.add_run().add_picture(img_path, width=Inches(0.85))
        else:
            r_c1_fail = p_c1.add_run("N/A")
            set_font(r_c1_fail, "Calibri", 8.5, False, COLOR_STEEL_GRAY)
            
        # Col 2: Descripción
        c2 = data_row.cells[2]
        c2.width = widths_op1[2]
        set_cell_background(c2, bg_row_color)
        set_cell_valign(c2, "center")
        p_c2 = c2.paragraphs[0]
        set_para_format(p_c2, before_pt=2, after_pt=2)
        r_c2 = p_c2.add_run(item["desc_op1"])
        set_font(r_c2, "Calibri", 9.5, False, COLOR_TEXT_DARK)
        
        # Col 3: Colores / Tallas
        c3 = data_row.cells[3]
        c3.width = widths_op1[3]
        set_cell_background(c3, bg_row_color)
        set_cell_valign(c3, "center")
        p_c3 = c3.paragraphs[0]
        set_para_format(p_c3, before_pt=2, after_pt=2)
        r_c3_col = p_c3.add_run(f"Colores:\n{item['colors_op1']}\nTallas: {item['sizes_op1']}")
        set_font(r_c3_col, "Calibri", 9, False, COLOR_TEXT_DARK)
        
        # Col 4: Precio Unitario
        c4 = data_row.cells[4]
        c4.width = widths_op1[4]
        set_cell_background(c4, bg_row_color)
        set_cell_valign(c4, "center")
        p_c4 = c4.paragraphs[0]
        p_c4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_c4 = p_c4.add_run(f"${item['price_op1']:.2f}")
        set_font(r_c4, "Calibri", 10, True, COLOR_MIDNIGHT_NAVY)
        
        # Col 5: Total
        c5 = data_row.cells[5]
        c5.width = widths_op1[5]
        set_cell_background(c5, bg_row_color)
        set_cell_valign(c5, "center")
        p_c5 = c5.paragraphs[0]
        p_c5.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total_val = item['price_op1'] * 200
        r_c5 = p_c5.add_run(f"${total_val:,.2f}")
        set_font(r_c5, "Montserrat Medium", 9.5, True, COLOR_MIDNIGHT_NAVY)
        
        # Margin cell overrides
        for cell_to_pad in [c0, c1, c2, c3, c4, c5]:
            set_cell_margins(cell_to_pad, top=80, bottom=80, left=80, right=80)

    # ---------------------------------------------
    # SECTION 5: OPCIÓN 2 – PROVEEDOR ID-MU (Page 6)
    # ---------------------------------------------
    print("Creating Opción 2...")
    doc.add_page_break()
    add_heading_styled(doc, "4. ALTERNATIVA IDENTIFICADA: OPCIÓN 2 – PROVEEDOR ID-MU", level=1)
    
    p_op2_desc = doc.add_paragraph()
    set_para_format(p_op2_desc, before_pt=0, after_pt=6)
    r_op2_desc = p_op2_desc.add_run("Esta alternativa ofrece modelos de estructura y diseño muy similares a la Opción 1, pero con precios unitarios significativamente menores, permitiendo un gran ahorro de capital.")
    set_font(r_op2_desc, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    
    # Small stats panel
    table_op2_stats = doc.add_table(1, 3)
    table_op2_stats.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table_op2_stats)
    stats_data_op2 = [
        ("INVERSIÓN TOTAL", "$35,979.00", HEX_MIDNIGHT_NAVY),
        ("CANTIDAD TOTAL", "3,000 unidades", HEX_STEEL_GRAY),
        ("AHORRO BRUTO", "$21,945.00 (37.9%)", HEX_COPPER)
    ]
    for idx, (title, value, color) in enumerate(stats_data_op2):
        cell_stat = table_op2_stats.cell(0, idx)
        cell_stat.width = Inches(2.3)
        set_cell_background(cell_stat, HEX_WARM_CREAM)
        set_cell_margins(cell_stat, top=80, bottom=80, left=100, right=100)
        set_cell_borders(cell_stat, top="single", color=HEX_COPPER, size="4")
        p_stat_t = cell_stat.paragraphs[0]
        p_stat_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_stat_t = p_stat_t.add_run(title + "\n")
        set_font(r_stat_t, "Montserrat Medium", 8.5, False, COLOR_STEEL_GRAY)
        r_stat_v = p_stat_t.add_run(value)
        set_font(r_stat_v, "Montserrat SemiBold", 12, True, RGBColor(int(color[0:2],16), int(color[2:4],16), int(color[4:6],16)))
        
    doc.add_paragraph() # spacer
    
    # 15 Product Table for Op 2
    table_op2 = doc.add_table(16, 6)
    table_op2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_op2, color="CCCCCC", size="4")
    
    headers_op2 = ["Item", "Foto", "Descripción / Especificaciones", "Colores / Tallas", "Precio", "Total (200 u.)"]
    widths_op2 = [Inches(0.4), Inches(0.95), Inches(2.45), Inches(1.4), Inches(0.8), Inches(1.0)]
    
    # Header Row Style
    hdr_row_op2 = table_op2.rows[0]
    hdr_row_op2.height = Inches(0.4)
    hdr_row_op2.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for i, title in enumerate(headers_op2):
        cell_hdr = hdr_row_op2.cells[i]
        cell_hdr.width = widths_op2[i]
        set_cell_background(cell_hdr, HEX_MIDNIGHT_NAVY)
        set_cell_margins(cell_hdr, top=100, bottom=100, left=100, right=100)
        set_cell_valign(cell_hdr, "center")
        p_hdr = cell_hdr.paragraphs[0]
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0,1,4,5] else WD_ALIGN_PARAGRAPH.LEFT
        r_hdr = p_hdr.add_run(title)
        set_font(r_hdr, "Montserrat SemiBold", 9, True, COLOR_WHITE)
        
    # Data Rows
    for idx, item in enumerate(ITEMS_DATA):
        row_idx = idx + 1
        data_row = table_op2.rows[row_idx]
        data_row.height = Inches(0.95)
        data_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        
        bg_row_color = HEX_WHITE if idx % 2 == 0 else HEX_WARM_CREAM
        
        # Col 0: Item
        c0 = data_row.cells[0]
        c0.width = widths_op2[0]
        set_cell_background(c0, bg_row_color)
        set_cell_valign(c0, "center")
        p_c0 = c0.paragraphs[0]
        p_c0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c0 = p_c0.add_run(str(item["id"]))
        set_font(r_c0, "Montserrat SemiBold", 10, True, COLOR_MIDNIGHT_NAVY)
        
        # Col 1: Foto
        c1 = data_row.cells[1]
        c1.width = widths_op2[1]
        set_cell_background(c1, bg_row_color)
        set_cell_valign(c1, "center")
        p_c1 = c1.paragraphs[0]
        p_c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_path = os.path.join("fotos blazer", item["img_op2"])
        if os.path.exists(img_path):
            p_c1.add_run().add_picture(img_path, width=Inches(0.85))
        else:
            r_c1_fail = p_c1.add_run("N/A")
            set_font(r_c1_fail, "Calibri", 8.5, False, COLOR_STEEL_GRAY)
            
        # Col 2: Descripción
        c2 = data_row.cells[2]
        c2.width = widths_op2[2]
        set_cell_background(c2, bg_row_color)
        set_cell_valign(c2, "center")
        p_c2 = c2.paragraphs[0]
        set_para_format(p_c2, before_pt=2, after_pt=2)
        r_c2 = p_c2.add_run(item["desc_op2"])
        set_font(r_c2, "Calibri", 9.5, False, COLOR_TEXT_DARK)
        
        # Col 3: Colores / Tallas
        c3 = data_row.cells[3]
        c3.width = widths_op2[3]
        set_cell_background(c3, bg_row_color)
        set_cell_valign(c3, "center")
        p_c3 = c3.paragraphs[0]
        set_para_format(p_c3, before_pt=2, after_pt=2)
        r_c3_col = p_c3.add_run(f"Colores:\n{item['colors_op2']}\nTallas: {item['sizes_op2']}")
        set_font(r_c3_col, "Calibri", 9, False, COLOR_TEXT_DARK)
        
        # Col 4: Precio Unitario
        c4 = data_row.cells[4]
        c4.width = widths_op2[4]
        set_cell_background(c4, bg_row_color)
        set_cell_valign(c4, "center")
        p_c4 = c4.paragraphs[0]
        p_c4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_c4 = p_c4.add_run(f"${item['price_op2']:.2f}")
        set_font(r_c4, "Calibri", 10, True, COLOR_MIDNIGHT_NAVY)
        
        # Col 5: Total
        c5 = data_row.cells[5]
        c5.width = widths_op2[5]
        set_cell_background(c5, bg_row_color)
        set_cell_valign(c5, "center")
        p_c5 = c5.paragraphs[0]
        p_c5.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total_val = item['price_op2'] * 200
        r_c5 = p_c5.add_run(f"${total_val:,.2f}")
        set_font(r_c5, "Montserrat Medium", 9.5, True, COLOR_MIDNIGHT_NAVY)
        
        # Margin cell overrides
        for cell_to_pad in [c0, c1, c2, c3, c4, c5]:
            set_cell_margins(cell_to_pad, top=80, bottom=80, left=80, right=80)

    # ---------------------------------------------
    # SECTION 6: COMPARATIVO EJECUTIVO DE PROVEEDORES (Page 7)
    # ---------------------------------------------
    print("Creating Comparativo Ejecutivo...")
    doc.add_page_break()
    add_heading_styled(doc, "5. COMPARATIVO EJECUTIVO DE PROVEEDORES", level=1)
    
    p_comp_desc = doc.add_paragraph()
    set_para_format(p_comp_desc, before_pt=0, after_pt=6)
    r_comp_desc = p_comp_desc.add_run("A continuación se muestra el análisis comparativo detallado modelo por modelo, permitiendo identificar las brechas de precios y el ahorro potencial proyectado para una orden base de 200 unidades por referencia.")
    set_font(r_comp_desc, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    
    # Executive Comparison Table
    table_comp = doc.add_table(17, 7) # 15 items + headers + totals
    table_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_comp, color="CCCCCC", size="4")
    
    headers_comp = ["Item", "Modelo", "P. Unit Op. 1", "P. Unit Op. 2", "Diferencia ($)", "Ahorro (%)", "Ahorro (200 u.)"]
    widths_comp = [Inches(0.4), Inches(2.2), Inches(0.9), Inches(0.9), Inches(0.9), Inches(0.8), Inches(1.1)]
    
    hdr_comp_row = table_comp.rows[0]
    hdr_comp_row.height = Inches(0.4)
    hdr_comp_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for i, title in enumerate(headers_comp):
        cell_hdr = hdr_comp_row.cells[i]
        cell_hdr.width = widths_comp[i]
        set_cell_background(cell_hdr, HEX_MIDNIGHT_NAVY)
        set_cell_margins(cell_hdr, top=100, bottom=100, left=80, right=80)
        set_cell_valign(cell_hdr, "center")
        p_hdr = cell_hdr.paragraphs[0]
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 2, 3, 4, 5, 6] else WD_ALIGN_PARAGRAPH.LEFT
        r_hdr = p_hdr.add_run(title)
        set_font(r_hdr, "Montserrat SemiBold", 8.5, True, COLOR_WHITE)
        
    tot_op1 = 0
    tot_op2 = 0
    tot_saved = 0
    
    for idx, item in enumerate(ITEMS_DATA):
        row_idx = idx + 1
        data_row = table_comp.rows[row_idx]
        bg_row_color = HEX_WHITE if idx % 2 == 0 else HEX_WARM_CREAM
        
        diff_price = item["price_op1"] - item["price_op2"]
        pct_saved = (diff_price / item["price_op1"]) * 100 if item["price_op1"] > 0 else 0
        saved_200 = diff_price * 200
        
        tot_op1 += item["price_op1"] * 200
        tot_op2 += item["price_op2"] * 200
        tot_saved += saved_200
        
        # Col 0: ID
        c0 = data_row.cells[0]
        c0.width = widths_comp[0]
        set_cell_background(c0, bg_row_color)
        p_c0 = c0.paragraphs[0]
        p_c0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c0 = p_c0.add_run(str(item["id"]))
        set_font(r_c0, "Montserrat SemiBold", 9.5, True, COLOR_MIDNIGHT_NAVY)
        
        # Col 1: Modelo
        c1 = data_row.cells[1]
        c1.width = widths_comp[1]
        set_cell_background(c1, bg_row_color)
        p_c1 = c1.paragraphs[0]
        r_c1 = p_c1.add_run(item["name"])
        set_font(r_c1, "Calibri", 9.5, False, COLOR_TEXT_DARK)
        
        # Col 2: P. Unit Op 1
        c2 = data_row.cells[2]
        c2.width = widths_comp[2]
        set_cell_background(c2, bg_row_color)
        p_c2 = c2.paragraphs[0]
        p_c2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_c2 = p_c2.add_run(f"${item['price_op1']:.2f}")
        set_font(r_c2, "Calibri", 9.5, False, COLOR_TEXT_DARK)
        
        # Col 3: P. Unit Op 2
        c3 = data_row.cells[3]
        c3.width = widths_comp[3]
        set_cell_background(c3, bg_row_color)
        p_c3 = c3.paragraphs[0]
        p_c3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_c3 = p_c3.add_run(f"${item['price_op2']:.2f}")
        set_font(r_c3, "Calibri", 9.5, True, COLOR_MIDNIGHT_NAVY)
        
        # Col 4: Diferencia
        c4 = data_row.cells[4]
        c4.width = widths_comp[4]
        set_cell_background(c4, bg_row_color)
        p_c4 = c4.paragraphs[0]
        p_c4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_c4 = p_c4.add_run(f"${diff_price:.2f}" if diff_price >= 0 else f"-${abs(diff_price):.2f}")
        # Color red for negative savings (Item 14)
        c4_color = COLOR_COPPER if diff_price >= 0 else RGBColor(0xBA, 0x3C, 0x2A)
        set_font(r_c4, "Calibri", 9.5, True, c4_color)
        
        # Col 5: Pct
        c5 = data_row.cells[5]
        c5.width = widths_comp[5]
        set_cell_background(c5, bg_row_color)
        p_c5 = c5.paragraphs[0]
        p_c5.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_c5 = p_c5.add_run(f"{pct_saved:.1f}%" if diff_price >= 0 else f"-{abs(pct_saved):.1f}%")
        set_font(r_c5, "Calibri", 9.5, False, c4_color)
        
        # Col 6: Ahorro Total
        c6 = data_row.cells[6]
        c6.width = widths_comp[6]
        set_cell_background(c6, bg_row_color)
        p_c6 = c6.paragraphs[0]
        p_c6.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_c6 = p_c6.add_run(f"${saved_200:,.2f}" if saved_200 >= 0 else f"-${abs(saved_200):,.2f}")
        set_font(r_c6, "Montserrat Medium", 9.5, True, c4_color)
        
        # Margin cell overrides
        for cell_to_pad in [c0, c1, c2, c3, c4, c5, c6]:
            set_cell_margins(cell_to_pad, top=80, bottom=80, left=80, right=80)
            set_cell_valign(cell_to_pad, "center")
            
    # Totals Row
    tot_row = table_comp.rows[16]
    tot_row.height = Inches(0.35)
    tot_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    # Manual merge for 'TOTALES'
    cell_tot_lbl = tot_row.cells[0]
    cell_tot_lbl.width = widths_comp[0]
    set_cell_background(cell_tot_lbl, HEX_MIDNIGHT_NAVY)
    p_tot_lbl = cell_tot_lbl.paragraphs[0]
    p_tot_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_tot_lbl = p_tot_lbl.add_run("TOTALES")
    set_font(r_tot_lbl, "Montserrat SemiBold", 8.5, True, COLOR_WHITE)
    
    # We can write 'Colección Completa' or empty space in next cell
    cell_tot_lbl2 = tot_row.cells[1]
    cell_tot_lbl2.width = widths_comp[1]
    set_cell_background(cell_tot_lbl2, HEX_MIDNIGHT_NAVY)
    p_tot_lbl2 = cell_tot_lbl2.paragraphs[0]
    r_tot_lbl2 = p_tot_lbl2.add_run("Colección Completa (3,000 unidades)")
    set_font(r_tot_lbl2, "Montserrat Medium", 8.5, False, COLOR_WHITE)
    
    cell_tot_op1 = tot_row.cells[2]
    cell_tot_op1.width = widths_comp[2]
    set_cell_background(cell_tot_op1, HEX_MIDNIGHT_NAVY)
    p_tot_op1 = cell_tot_op1.paragraphs[0]
    p_tot_op1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_tot_op1 = p_tot_op1.add_run(f"${tot_op1:,.2f}")
    set_font(r_tot_op1, "Montserrat SemiBold", 9, True, COLOR_WHITE)
    
    cell_tot_op2 = tot_row.cells[3]
    cell_tot_op2.width = widths_comp[3]
    set_cell_background(cell_tot_op2, HEX_MIDNIGHT_NAVY)
    p_tot_op2 = cell_tot_op2.paragraphs[0]
    p_tot_op2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_tot_op2 = p_tot_op2.add_run(f"${tot_op2:,.2f}")
    set_font(r_tot_op2, "Montserrat SemiBold", 9, True, COLOR_WHITE)
    
    # Empty diff
    cell_tot_diff = tot_row.cells[4]
    cell_tot_diff.width = widths_comp[4]
    set_cell_background(cell_tot_diff, HEX_MIDNIGHT_NAVY)
    
    # Total Pct Ahorro
    overall_pct_saving = (tot_saved / tot_op1) * 100
    cell_tot_pct = tot_row.cells[5]
    cell_tot_pct.width = widths_comp[5]
    set_cell_background(cell_tot_pct, HEX_MIDNIGHT_NAVY)
    p_tot_pct = cell_tot_pct.paragraphs[0]
    p_tot_pct.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_tot_pct = p_tot_pct.add_run(f"{overall_pct_saving:.1f}%")
    set_font(r_tot_pct, "Montserrat SemiBold", 9, True, COLOR_GOLD_LIGHT)
    
    # Total Ahorro
    cell_tot_saved = tot_row.cells[6]
    cell_tot_saved.width = widths_comp[6]
    set_cell_background(cell_tot_saved, HEX_MIDNIGHT_NAVY)
    p_tot_saved = cell_tot_saved.paragraphs[0]
    p_tot_saved.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_tot_saved = p_tot_saved.add_run(f"${tot_saved:,.2f}")
    set_font(r_tot_saved, "Montserrat SemiBold", 9.5, True, COLOR_GOLD_LIGHT)
    
    for cell_to_pad in [cell_tot_lbl, cell_tot_lbl2, cell_tot_op1, cell_tot_op2, cell_tot_diff, cell_tot_pct, cell_tot_saved]:
        set_cell_margins(cell_to_pad, top=100, bottom=100, left=80, right=80)
        set_cell_valign(cell_to_pad, "center")

    doc.add_paragraph() # spacer
    
    p_comp_find = doc.add_paragraph()
    set_para_format(p_comp_find, before_pt=12, after_pt=4)
    r_comp_find = p_comp_find.add_run("Principales Hallazgos y Análisis del Comparativo:")
    set_font(r_comp_find, "Montserrat SemiBold", 11, True, COLOR_MIDNIGHT_NAVY)
    
    add_bullet_point(doc, "Diferencias críticas de precios", "La gran mayoría de referencias presenta brechas de precio superiores al 30%, con casos extremos como la Chaqueta Chanel (Item 4, 67.8% más barata) y el Blazer Moderno (Item 8, 80.5% más barato) en la Opción 2.")
    add_bullet_point(doc, "Impacto presupuestario", "El costo total de importar toda la colección con la Opción 1 asciende a USD 57,924.00, mientras que con la Opción 2 se reduce a USD 35,979.00. Esto representa un ahorro directo de USD 21,945.00 para la misma cantidad de unidades.")
    add_bullet_point(doc, "El factor Chaqueta Biker", "La Chaqueta Biker (Item 14) es el único producto donde la Opción 1 es más económica ($32.40 vs $52.63). Esto se debe a diferencias de materiales y estructura. Por ende, no se recomienda adquirir este artículo con la Opción 2.")
    add_bullet_point(doc, "Viabilidad del piloto comercial", "La Opción 2 permite construir una colección piloto con menor riesgo, facilitando la validación del mercado antes de expandir el volumen de importación.")

    # ---------------------------------------------
    # SECTION 7: MODELOS DESTACADOS (Page 8 & 9)
    # ---------------------------------------------
    print("Creating Modelos Destacados...")
    doc.add_page_break()
    add_heading_styled(doc, "6. MODELOS DESTACADOS CON ALTO POTENCIAL COMERCIAL", level=1)
    
    p_dest_intro = doc.add_paragraph()
    set_para_format(p_dest_intro, before_pt=0, after_pt=12)
    r_dest_intro = p_dest_intro.add_run("Hemos seleccionado cinco prendas clave de la colección que destacan por su diseño sofisticado, alta percepción de valor y excelente proyección de rotación comercial para una primera importación piloto. Para una presentación premium, se analiza detalladamente cada modelo:")
    set_font(r_dest_intro, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    
    featured_models = [
        {
            "num": 1,
            "title": "Blazer blanco doble botonadura (Modelo Destacado 1)",
            "item_id": 1,
            "img": "1.jpeg",
            "advantages": [
                "Alta percepción de valor gracias a su sastrero clásico estructurado.",
                "Corte con doble botonadura dorada de gran tendencia en el mercado corporativo y casual elegante.",
                "Estructura premium que realza la silueta, ideal como prenda ancla de la colección."
            ],
            "potencial": "Muy Alto (Segmento ejecutivo/casual premium).",
            "comentario": "El blazer blanco es un básico atemporal de gran demanda todo el año. Su alta percepción de valor justifica plenamente su incorporación en el portafolio piloto."
        },
        {
            "num": 2,
            "title": "Blazer rosado tweed (Modelo Destacado 2)",
            "item_id": 2,
            "img": "2.jpeg",
            "advantages": [
                "Diseño contemporáneo en tejido de tweed bouclé rosa, altamente demandado en la moda femenina.",
                "Corte princesa estructurado que garantiza un ajuste excelente y estilizado.",
                "Gran versatilidad comercial para campañas de media estación y temporadas de invierno/primavera."
            ],
            "potencial": "Excelente (Segmento juvenil y contemporáneo).",
            "comentario": "La textura de tweed bouclé en tonos rosa es tendencia internacional. Genera un gran atractivo visual y diferenciación en vitrina o catálogo digital."
        },
        {
            "num": 3,
            "title": "Chaqueta negra estilo Chanel (Modelo Destacado 3)",
            "item_id": 4,
            "img": "4.jpeg",
            "advantages": [
                "Corte clásico cropped estilo Chanel, una pieza icónica y atemporal de la moda femenina.",
                "El color negro facilita una altísima rotación comercial y reduce el riesgo de inventario lento.",
                "Precio de adquisición sumamente competitivo en la Opción 2 ($5.79), maximizando el margen de ganancia."
            ],
            "potencial": "Extraordinario (Rotación garantizada por ser un básico atemporal).",
            "comentario": "Es el modelo más rentable de la evaluación. Combina una demanda masiva, bajo riesgo comercial y un costo de importación excepcionalmente bajo."
        },
        {
            "num": 4,
            "title": "Abrigo camel largo (Modelo Destacado 4)",
            "item_id": 6,
            "img": "6.jpeg",
            "advantages": [
                "Prenda exterior de ticket promedio superior, ideal para incrementar la facturación global.",
                "Corte princesa con solapa clásica y botones dorados que aportan elegancia otoñal.",
                "Colores comerciales de alta demanda invernal (camel, caqui y negro)."
            ],
            "potencial": "Alto (Temporada otoño-invierno, alto ticket de venta).",
            "comentario": "Representa la prenda de mayor valor de la colección. Su inclusión en el catálogo aporta madurez visual y eleva la imagen de marca de Importaciones Manorca S.A.C."
        },
        {
            "num": 5,
            "title": "Cárdigan blanco (Modelo Destacado 5)",
            "item_id": 10,
            "img": "10.jpeg",
            "advantages": [
                "Textura tejida de punto suave con ribetes negros en contraste, logrando un estilo minimalista elegante.",
                "Prenda cómoda y ligera, ideal para media estación y uso diario.",
                "Precio de adquisición altamente accesible ($7.07), facilitando compras complementarias o de impulso."
            ],
            "potencial": "Medio-Alto (Rotación constante para uso diario).",
            "comentario": "Funciona perfectamente como producto complementario. Atrae a clientes que buscan una opción más relajada pero que mantenga un estándar elegante."
        }
    ]
    
    # Render models in a clean layout (1x2 table for each model: Left photo, Right details)
    for idx, model in enumerate(featured_models):
        if idx == 2:  # Page break after first 2 models to balance pages 8 and 9!
            doc.add_page_break()
            add_heading_styled(doc, "6. MODELOS DESTACADOS (CONTINUACIÓN)", level=1)
            
        p_mod_title = doc.add_paragraph()
        set_para_format(p_mod_title, before_pt=12, after_pt=6, keep_with_next=True)
        r_mod_num = p_mod_title.add_run(f"N° {model['num']}: ")
        set_font(r_mod_num, "Montserrat SemiBold", 12, True, COLOR_COPPER)
        r_mod_t = p_mod_title.add_run(model["title"])
        set_font(r_mod_t, "Montserrat SemiBold", 12, True, COLOR_MIDNIGHT_NAVY)
        
        mod_table = doc.add_table(1, 2)
        mod_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        mod_table.autofit = False
        remove_table_borders(mod_table)
        
        # Left Cell: Image
        cell_img = mod_table.cell(0, 0)
        cell_img.width = Inches(2.2)
        set_cell_valign(cell_img, "top")
        set_cell_background(cell_img, HEX_WHITE)
        set_cell_margins(cell_img, top=20, bottom=20, left=20, right=80)
        
        p_img = cell_img.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_path = os.path.join("fotos blazer", model["img"])
        if os.path.exists(img_path):
            p_img.add_run().add_picture(img_path, width=Inches(2.1))
        else:
            r_img_fail = p_img.add_run("[IMAGEN NO DISPONIBLE]")
            set_font(r_img_fail, "Calibri", 9, False, COLOR_STEEL_GRAY)
            
        # Right Cell: Detailed Info
        cell_info = mod_table.cell(0, 1)
        cell_info.width = Inches(4.8)
        set_cell_valign(cell_info, "top")
        set_cell_background(cell_info, HEX_WHITE)
        set_cell_margins(cell_info, top=20, bottom=20, left=80, right=20)
        
        p_adv_lbl = cell_info.paragraphs[0]
        set_para_format(p_adv_lbl, before_pt=0, after_pt=4)
        r_adv_lbl = p_adv_lbl.add_run("Ventajas Clave:")
        set_font(r_adv_lbl, "Montserrat Medium", 10, True, COLOR_COPPER)
        
        for adv in model["advantages"]:
            p_adv = cell_info.add_paragraph()
            set_para_format(p_adv, before_pt=1, after_pt=2)
            p_adv.paragraph_format.left_indent = Inches(0.15)
            r_bullet = p_adv.add_run("▪  ")
            set_font(r_bullet, "Calibri", 9, True, COLOR_STEEL_GRAY)
            r_text = p_adv.add_run(adv)
            set_font(r_text, "Calibri", 9.5, False, COLOR_TEXT_DARK)
            
        p_pot = cell_info.add_paragraph()
        set_para_format(p_pot, before_pt=4, after_pt=2)
        r_pot_lbl = p_pot.add_run("Potencial Comercial: ")
        set_font(r_pot_lbl, "Montserrat Medium", 9.5, True, COLOR_MIDNIGHT_NAVY)
        r_pot_val = p_pot.add_run(model["potencial"])
        set_font(r_pot_val, "Calibri", 9.5, False, COLOR_TEXT_DARK)
        
        p_com = cell_info.add_paragraph()
        set_para_format(p_com, before_pt=4, after_pt=0)
        p_com.paragraph_format.left_indent = Inches(0.15)
        # We put comments in a subtle light gray card style
        set_cell_margins(cell_info, top=20, bottom=20, left=80, right=20)
        r_com_lbl = p_com.add_run("Smart Import Trade Comment: ")
        set_font(r_com_lbl, "Montserrat SemiBold", 9, True, COLOR_COPPER)
        r_com_val = p_com.add_run(model["comentario"])
        set_font(r_com_val, "Calibri", 9.5, True, COLOR_MIDNIGHT_NAVY)
        
        doc.add_paragraph() # spacing between models

    # ---------------------------------------------
    # SECTION 8: RANKING SMART IMPORT TRADE (Page 10)
    # ---------------------------------------------
    print("Creating Ranking Page...")
    doc.add_page_break()
    add_heading_styled(doc, "7. RANKING SMART IMPORT TRADE", level=1)
    
    p_rank_desc = doc.add_paragraph()
    set_para_format(p_rank_desc, before_pt=0, after_pt=12)
    r_rank_desc = p_rank_desc.add_run("Para orientar la selección definitiva de modelos para la importación piloto, hemos elaborado un ranking cuantitativo ponderando tres criterios clave en escala de 1 a 10: ")
    set_font(r_rank_desc, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    r_rank_c = p_rank_desc.add_run("Precio (estructura de costo y rentabilidad), Diseño (fidelidad y atractivo estético) y Potencial Comercial (rotación y versatilidad).")
    set_font(r_rank_c, "Calibri", 10.5, True, COLOR_MIDNIGHT_NAVY)
    
    # Ranking Table
    table_rank = doc.add_table(6, 6)
    table_rank.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_rank, color="CCCCCC", size="4")
    
    headers_rank = ["Puesto", "Modelo", "Costo", "Diseño", "Rotación", "Puntaje Total"]
    widths_rank = [Inches(0.8), Inches(2.4), Inches(0.9), Inches(0.9), Inches(0.9), Inches(1.1)]
    
    hdr_rank_row = table_rank.rows[0]
    hdr_rank_row.height = Inches(0.4)
    hdr_rank_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for i, title in enumerate(headers_rank):
        cell_hdr = hdr_rank_row.cells[i]
        cell_hdr.width = widths_rank[i]
        set_cell_background(cell_hdr, HEX_MIDNIGHT_NAVY)
        set_cell_margins(cell_hdr, top=100, bottom=100, left=80, right=80)
        set_cell_valign(cell_hdr, "center")
        p_hdr = cell_hdr.paragraphs[0]
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_hdr = p_hdr.add_run(title)
        set_font(r_hdr, "Montserrat SemiBold", 9, True, COLOR_WHITE)
        
    ranking_data = [
        ("🥇 1° Puesto", "Chaqueta Chanel negra (Item 4)", "10/10", "8/10", "9/10", "36/40", "Excelente relación costo/beneficio. El modelo más rentable de toda la colección."),
        ("🥈 2° Puesto", "Cárdigan punto texturizado (Item 10)", "8/10", "8/10", "9/10", "34/40", "Atemporal, de bajo costo y alta demanda en media estación."),
        ("🥉 3° Puesto", "Cárdigan punto acanalado (Item 11)", "10/10", "7/10", "8/10", "34/40", "Excelente precio en Opción 2 ($4.52), ideal para compras de impulso."),
        ("4° Lugar", "Abrigo camel largo (Item 6)", "8/10", "9/10", "8/10", "33/40", "Aporta alta gama a la colección, ideal para venta de alto ticket."),
        ("5° Lugar", "Blazer blanco doble botonadura (Item 1)", "6/10", "9/10", "9/10", "32/40", "Diseño sofisticado pero con un precio unitario mayor respecto a los cárdigans.")
    ]
    
    for idx, (puesto, modelo, costo, dis, rot, total, just) in enumerate(ranking_data):
        row_idx = idx + 1
        data_row = table_rank.rows[row_idx]
        data_row.height = Inches(0.45)
        data_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        bg_row_color = HEX_WHITE if idx % 2 == 0 else HEX_WARM_CREAM
        
        # Puesto
        c0 = data_row.cells[0]
        c0.width = widths_rank[0]
        set_cell_background(c0, bg_row_color)
        p_c0 = c0.paragraphs[0]
        p_c0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c0 = p_c0.add_run(puesto)
        set_font(r_c0, "Montserrat SemiBold", 9.5, True, COLOR_COPPER if idx < 3 else COLOR_STEEL_GRAY)
        
        # Modelo
        c1 = data_row.cells[1]
        c1.width = widths_rank[1]
        set_cell_background(c1, bg_row_color)
        p_c1 = c1.paragraphs[0]
        r_c1 = p_c1.add_run(modelo)
        set_font(r_c1, "Calibri", 9.5, True, COLOR_MIDNIGHT_NAVY)
        
        # Costo
        c2 = data_row.cells[2]
        c2.width = widths_rank[2]
        set_cell_background(c2, bg_row_color)
        p_c2 = c2.paragraphs[0]
        p_c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c2 = p_c2.add_run(costo)
        set_font(r_c2, "Calibri", 9.5, False, COLOR_TEXT_DARK)
        
        # Diseño
        c3 = data_row.cells[3]
        c3.width = widths_rank[3]
        set_cell_background(c3, bg_row_color)
        p_c3 = c3.paragraphs[0]
        p_c3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c3 = p_c3.add_run(dis)
        set_font(r_c3, "Calibri", 9.5, False, COLOR_TEXT_DARK)
        
        # Rotación
        c4 = data_row.cells[4]
        c4.width = widths_rank[4]
        set_cell_background(c4, bg_row_color)
        p_c4 = c4.paragraphs[0]
        p_c4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c4 = p_c4.add_run(rot)
        set_font(r_c4, "Calibri", 9.5, False, COLOR_TEXT_DARK)
        
        # Total
        c5 = data_row.cells[5]
        c5.width = widths_rank[5]
        set_cell_background(c5, bg_row_color)
        p_c5 = c5.paragraphs[0]
        p_c5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c5 = p_c5.add_run(total)
        set_font(r_c5, "Montserrat SemiBold", 10, True, COLOR_MIDNIGHT_NAVY)
        
        for cell_to_pad in [c0, c1, c2, c3, c4, c5]:
            set_cell_margins(cell_to_pad, top=80, bottom=80, left=80, right=80)
            set_cell_valign(cell_to_pad, "center")

    doc.add_paragraph() # spacer
    
    p_just_lbl = doc.add_paragraph()
    set_para_format(p_just_lbl, before_pt=12, after_pt=4, keep_with_next=True)
    r_just_lbl = p_just_lbl.add_run("Justificación Detallada del Top 3:")
    set_font(r_just_lbl, "Montserrat SemiBold", 11, True, COLOR_MIDNIGHT_NAVY)
    
    for idx, (puesto, modelo, _, _, _, _, just) in enumerate(ranking_data[:3]):
        p_just = doc.add_paragraph()
        set_para_format(p_just, before_pt=2, after_pt=4)
        p_just.paragraph_format.left_indent = Inches(0.2)
        r_bullet = p_just.add_run(puesto.split()[0] + "  ")
        set_font(r_bullet, "Calibri", 10.5, True, COLOR_COPPER)
        r_mod_name = p_just.add_run(modelo + ": ")
        set_font(r_mod_name, "Montserrat Medium", 10, True, COLOR_MIDNIGHT_NAVY)
        r_just_body = p_just.add_run(just)
        set_font(r_just_body, "Calibri", 10.5, False, COLOR_TEXT_DARK)

    # ---------------------------------------------
    # SECTION 9: PRODUCTOS COMPLEMENTARIOS (Page 11)
    # ---------------------------------------------
    print("Creating Productos Complementarios...")
    doc.add_page_break()
    add_heading_styled(doc, "8. ANÁLISIS DE PRODUCTOS COMPLEMENTARIOS", level=1)
    
    p_compl_desc = doc.add_paragraph()
    set_para_format(p_compl_desc, before_pt=0, after_pt=12)
    r_compl_desc = p_compl_desc.add_run("Durante la búsqueda de proveedores, se identificaron artículos complementarios que no forman parte de la colección de sastrería o abrigo principal, pero que ofrecen una excelente oportunidad de negocio lateral o de impulso comercial:")
    set_font(r_compl_desc, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    
    compl_items = [
        {
            "name": "Blusa entallada de gasa (Item 13)",
            "desc": "Blusa entallada de manga larga con detalles de volantes y botones. Excelente combinación para acompañar cualquiera de los blazers de tweed.",
            "op1": "$10.80",
            "op2": "$3.20",
            "img": "13opcion 2.jpg",
            "comment": "Prenda con excelente costo en Opción 2. Ideal para ofrecer el look completo a la cliente final."
        },
        {
            "name": "Chaqueta Biker de cuero sintético (Item 14)",
            "desc": "Chaqueta biker de PU estilo Balmain con botones metálicos decorativos. Aporta un perfil juvenil y rebelde.",
            "op1": "$32.40",
            "op2": "$52.63",
            "img": "14.jpeg",
            "comment": "¡Alerta de costo! La Opción 1 es sustancialmente más económica aquí. Evitar comprar a la Opción 2."
        },
        {
            "name": "Termo de acero inoxidable 40 oz (Item 15)",
            "desc": "Termo térmico de 40 oz con aislamiento al vacío, asa de sujeción y tapa hermética. Colores pastel de alta tendencia.",
            "op1": "$14.40",
            "op2": "$8.44",
            "img": "15.jpeg",
            "comment": "Artículo de merchandising con alta rentabilidad para campañas de fidelización o venta cruzada."
        }
    ]
    
    for idx, c_item in enumerate(compl_items):
        p_c_title = doc.add_paragraph()
        set_para_format(p_c_title, before_pt=12, after_pt=6, keep_with_next=True)
        r_c_title = p_c_title.add_run(c_item["name"])
        set_font(r_c_title, "Montserrat SemiBold", 11.5, True, COLOR_MIDNIGHT_NAVY)
        
        c_table = doc.add_table(1, 2)
        c_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        c_table.autofit = False
        remove_table_borders(c_table)
        
        # Left: Image
        cell_img = c_table.cell(0, 0)
        cell_img.width = Inches(2.0)
        set_cell_valign(cell_img, "center")
        set_cell_background(cell_img, HEX_WHITE)
        set_cell_margins(cell_img, top=20, bottom=20, left=20, right=80)
        
        p_img = cell_img.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_path = os.path.join("fotos blazer", c_item["img"])
        if os.path.exists(img_path):
            p_img.add_run().add_picture(img_path, width=Inches(1.8))
        else:
            r_img_fail = p_img.add_run("[IMAGEN]")
            set_font(r_img_fail, "Calibri", 9, False, COLOR_STEEL_GRAY)
            
        # Right: Details
        cell_info = c_table.cell(0, 1)
        cell_info.width = Inches(5.0)
        set_cell_valign(cell_info, "top")
        set_cell_background(cell_info, HEX_WHITE)
        set_cell_margins(cell_info, top=20, bottom=20, left=80, right=20)
        
        p_desc = cell_info.paragraphs[0]
        set_para_format(p_desc, before_pt=0, after_pt=4)
        r_desc = p_desc.add_run(c_item["desc"])
        set_font(r_desc, "Calibri", 9.5, False, COLOR_TEXT_DARK)
        
        p_costs = cell_info.add_paragraph()
        set_para_format(p_costs, before_pt=4, after_pt=4)
        r_cost_lbl = p_costs.add_run("Comparativo de Costos Unitarios:\n")
        set_font(r_cost_lbl, "Montserrat Medium", 9.5, True, COLOR_MIDNIGHT_NAVY)
        r_cost_val = p_costs.add_run(f"  •  Costo Opción 1: {c_item['op1']}\n  •  Costo Opción 2: {c_item['op2']}")
        set_font(r_cost_val, "Calibri", 9.5, False, COLOR_TEXT_DARK)
        
        p_com = cell_info.add_paragraph()
        set_para_format(p_com, before_pt=4, after_pt=0)
        r_com_lbl = p_com.add_run("Observación: ")
        set_font(r_com_lbl, "Montserrat SemiBold", 9, True, COLOR_COPPER)
        r_com_val = p_com.add_run(c_item["comment"])
        set_font(r_com_val, "Calibri", 9.5, True, COLOR_MIDNIGHT_NAVY)
        
        doc.add_paragraph() # spacing

    # ---------------------------------------------
    # SECTION 10: RECOMENDACIÓN ESTRATÉGICA (Page 12)
    # ---------------------------------------------
    print("Creating Recomendación Estratégica...")
    doc.add_page_break()
    add_heading_styled(doc, "9. RECOMENDACIÓN ESTRATÉGICA DE INVERSIÓN", level=1)
    
    p_reco_intro = doc.add_paragraph()
    set_para_format(p_reco_intro, before_pt=0, after_pt=12)
    r_reco_intro = p_reco_intro.add_run("Con un presupuesto estimado disponible de aproximadamente ")
    set_font(r_reco_intro, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    r_reco_b = p_reco_intro.add_run("USD 5,000.00")
    set_font(r_reco_b, "Calibri", 10.5, True, COLOR_COPPER)
    r_reco_end = p_reco_intro.add_run(", Smart Import Trade recomienda seguir una estrategia selectiva de compras piloto enfocada en optimizar la inversión y reducir el riesgo comercial:")
    set_font(r_reco_end, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    
    create_callout_box(
        doc,
        "ESTRATEGIA RECOMENDADA DE COMPRA PILOTO (PRESUPUESTO USD 5,000)",
        "1. No incorporar la totalidad de las 15 referencias en el primer embarque, ya que cada referencia requiere un MOQ de 200 unidades, lo cual excedería ampliamente el presupuesto.\n\n"
        "2. Seleccionar una colección cápsula de entre 4 y 5 modelos con mayor potencial comercial de la Opción 2.\n\n"
        "3. Priorizar los siguientes modelos que juntos permiten optimizar el presupuesto disponible:\n"
        "  • Item 4 (Chaqueta Chanel negra - 200 unidades): USD 1,158.00\n"
        "  • Item 11 (Cárdigan punto acanalado - 200 unidades): USD 903.00\n"
        "  • Item 10 (Cárdigan punto texturizado - 200 unidades): USD 1,414.00\n"
        "  • Item 6 (Abrigo camel largo - 200 unidades): USD 1,649.00\n"
        "  ➔ INVERSIÓN TOTAL ESTIMADA: USD 5,124.00 (800 prendas listas para importar).\n\n"
        "4. En la Opción 1, este mismo mix piloto representaría un costo de USD 10,500.00 (una diferencia del 104% en contra de la rentabilidad del cliente)."
    )
    
    p_reco_fin = doc.add_paragraph()
    set_para_format(p_reco_fin, before_pt=6, after_pt=6)
    r_reco_fin = p_reco_fin.add_run("Esta distribución permite a Importaciones Manorca S.A.C. salir al mercado peruano con una colección representativa de blazers, cárdigans y abrigos sin sobrepasar su capacidad financiera inicial, maximizando los márgenes comerciales.")
    set_font(r_reco_fin, "Calibri", 10.5, False, COLOR_TEXT_DARK)

    # ---------------------------------------------
    # SECTION 11: VALIDACIÓN DE PRODUCTOS (Page 13)
    # ---------------------------------------------
    print("Creating Validación de Productos...")
    doc.add_page_break()
    add_heading_styled(doc, "10. PROCEDIMIENTO DE VALIDACIÓN DE PRODUCTOS", level=1)
    
    p_val_desc = doc.add_paragraph()
    set_para_format(p_val_desc, before_pt=0, after_pt=12)
    r_val_desc = p_val_desc.add_run("Antes de formalizar cualquier depósito o confirmar la orden de compra definitiva con el fabricante seleccionado, Smart Import Trade establece de forma obligatoria el siguiente proceso de control y validación de calidad:")
    set_font(r_val_desc, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    
    add_bullet_point(doc, "Solicitud de Fotografías Reales", "El proveedor debe suministrar fotos reales de alta definición de las telas y costuras en primer plano, sin filtros ni retoques digitales, correspondientes a los modelos seleccionados.")
    add_bullet_point(doc, "Solicitud de Videos Reales", "Videos cortos del taller de confección y de las prendas en movimiento sobre maniquí para constatar la caída del tejido, el brillo de los botones dorados y la flexibilidad del punto.")
    add_bullet_point(doc, "Confirmación Ficha Técnica de Materiales", "Declaración jurada de la composición textil (porcentaje exacto de poliéster, acrílico, lana o algodón) para evitar problemas en el etiquetado aduanero y la posterior nacionalización en Perú.")
    add_bullet_point(doc, "Validación de Cuadro de Tallas y Colores", "Confirmación de las medidas exactas por talla (S, M, L) expresadas en centímetros adaptadas al biotipo latino, y la disponibilidad exacta de stock en los colores solicitados.")
    add_bullet_point(doc, "Desarrollo de Muestras Físicas Opcionales", "En caso de requerirse una seguridad total sobre la calidad del bouclé o tweed, se programará el envío de muestras por courier expreso (costo asumido por el cliente, reembolsable al emitir la orden principal).")

    # ---------------------------------------------
    # SECTION 12: PRÓXIMOS PASOS (Page 14)
    # ---------------------------------------------
    print("Creating Próximos Pasos...")
    doc.add_page_break()
    add_heading_styled(doc, "11. CRONOGRAMA DE PRÓXIMOS PASOS", level=1)
    
    p_time_desc = doc.add_paragraph()
    set_para_format(p_time_desc, before_pt=0, after_pt=12)
    r_time_desc = p_time_desc.add_run("Para avanzar ordenadamente hacia la importación de la colección, proponemos el siguiente cronograma de trabajo con hitos definidos:")
    set_font(r_time_desc, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    
    # Timeline Table Layout (6 blocks in a table, styled elegantly)
    table_time = doc.add_table(1, 6)
    table_time.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_time.autofit = False
    remove_table_borders(table_time)
    
    steps = [
        ("1. Selección", "Definición final de los 4 o 5 modelos del catálogo.", HEX_MIDNIGHT_NAVY),
        ("2. Validación", "Solicitud de fotos y videos reales al fabricante.", HEX_COPPER),
        ("3. Muestras", "Envío e inspección física de muestras (opcional).", HEX_GOLD_LIGHT),
        ("4. Costeo", "Cálculo logístico y tributario integral puesto en Perú.", HEX_MIDNIGHT_NAVY),
        ("5. Viabilidad", "Evaluación final de márgenes y rentabilidad comercial.", HEX_COPPER),
        ("6. Orden", "Formalización de la orden de compra y pago inicial.", HEX_GOLD_LIGHT)
    ]
    
    for idx, (step_t, step_d, color) in enumerate(steps):
        cell_step = table_time.cell(0, idx)
        cell_step.width = Inches(1.16)
        set_cell_background(cell_step, HEX_WARM_CREAM)
        set_cell_margins(cell_step, top=100, bottom=100, left=80, right=80)
        set_cell_valign(cell_step, "top")
        
        # Border top color matching corporate colors
        set_cell_borders(cell_step, top="single", color=color, size="24")
        
        p_st = cell_step.paragraphs[0]
        set_para_format(p_st, before_pt=0, after_pt=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        r_st = p_st.add_run(step_t)
        set_font(r_st, "Montserrat SemiBold", 8.5, True, RGBColor(int(color[0:2],16), int(color[2:4],16), int(color[4:6],16)))
        
        p_sd = cell_step.add_paragraph()
        set_para_format(p_sd, before_pt=0, after_pt=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        r_sd = p_sd.add_run(step_d)
        set_font(r_sd, "Calibri", 8, False, COLOR_TEXT_DARK)
        
    doc.add_paragraph() # spacer
    
    # Adding arrow diagram text
    p_diagram = doc.add_paragraph()
    set_para_format(p_diagram, before_pt=12, after_pt=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    r_diag = p_diagram.add_run(
        "Selección de Modelos  ➔  Fotos y Videos Reales  ➔  Muestras Físicas (Opcional)\n"
        "➔  Costeo Integral Aduanero  ➔  Evaluación de Viabilidad Final  ➔  Orden de Compra Definitiva"
    )
    set_font(r_diag, "Montserrat Medium", 9.5, True, COLOR_COPPER)

    # ---------------------------------------------
    # SECTION 13: CONCLUSIÓN EJECUTIVA (Page 15)
    # ---------------------------------------------
    print("Creating Conclusión Ejecutiva...")
    doc.add_page_break()
    add_heading_styled(doc, "12. CONCLUSIÓN EJECUTIVA Y PRÓXIMOS PASOS", level=1)
    
    p_c1 = doc.add_paragraph()
    set_para_format(p_c1, before_pt=0, after_pt=6)
    r_c1 = p_c1.add_run("La investigación realizada por Smart Import Trade confirma que el desarrollo de la colección de blazers, tweed y cárdigans para ")
    set_font(r_c1, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    r_c1_b = p_c1.add_run("Importaciones Manorca S.A.C.")
    set_font(r_c1_b, "Calibri", 10.5, True, COLOR_MIDNIGHT_NAVY)
    r_c1_e = p_c1.add_run(" es técnica y comercialmente viable. La existencia de la Opción 2 permite ingresar al mercado con un producto altamente competitivo, minimizando la inversión y asegurando márgenes atractivos de rentabilidad.")
    set_font(r_c1_e, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    
    p_c2 = doc.add_paragraph()
    set_para_format(p_c2, before_pt=0, after_pt=18)
    r_c2 = p_c2.add_run("Se recomienda avanzar inmediatamente hacia la selección definitiva de modelos del mix de prueba sugerido, la posterior validación visual y técnica mediante el envío de fotos y videos reales por parte del proveedor preseleccionado, y la correspondiente cotización del flete internacional y aranceles para tener el costeo integral en almacén Lima.")
    set_font(r_c2, "Calibri", 10.5, False, COLOR_TEXT_DARK)
    
    p_sig = doc.add_paragraph()
    set_para_format(p_sig, before_pt=24, after_pt=6, keep_with_next=True)
    r_sig = p_sig.add_run("Firmas y Aprobaciones del Informe:")
    set_font(r_sig, "Montserrat Medium", 11, True, COLOR_MIDNIGHT_NAVY)
    
    sig_table = doc.add_table(1, 2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(sig_table)
    
    # Cell 1: Smart Import Trade
    cell_s1 = sig_table.cell(0, 0)
    cell_s1.width = Inches(3.5)
    set_cell_background(cell_s1, HEX_WHITE)
    set_cell_margins(cell_s1, top=100, bottom=100, left=100, right=100)
    p_s1 = cell_s1.paragraphs[0]
    p_s1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s1 = p_s1.add_run("\n\n___________________________________\n")
    set_font(r_s1, "Calibri", 11, False, COLOR_STEEL_GRAY)
    r_s1_name = p_s1.add_run("Smart Import Trade Sourcing Team\n")
    set_font(r_s1_name, "Montserrat Medium", 9.5, True, COLOR_MIDNIGHT_NAVY)
    r_s1_pos = p_s1.add_run("División de Comercio Exterior")
    set_font(r_s1_pos, "Calibri", 9, False, COLOR_STEEL_GRAY)
    
    # Cell 2: Importaciones Manorca
    cell_s2 = sig_table.cell(0, 1)
    cell_s2.width = Inches(3.5)
    set_cell_background(cell_s2, HEX_WHITE)
    set_cell_margins(cell_s2, top=100, bottom=100, left=100, right=100)
    p_s2 = cell_s2.paragraphs[0]
    p_s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s2 = p_s2.add_run("\n\n___________________________________\n")
    set_font(r_s2, "Calibri", 11, False, COLOR_STEEL_GRAY)
    r_s2_name = p_s2.add_run("Dirección General / Gerencia de Compras\n")
    set_font(r_s2_name, "Montserrat Medium", 9.5, True, COLOR_MIDNIGHT_NAVY)
    r_s2_pos = p_s2.add_run("Importaciones Manorca S.A.C.")
    set_font(r_s2_pos, "Calibri", 9, False, COLOR_STEEL_GRAY)

    # ---------------------------------------------
    # SECTION 14: CONTRAPORTADA (BACK COVER PAGE)
    # ---------------------------------------------
    print("Creating Back Cover Page...")
    section_back = doc.add_section(docx.enum.section.WD_SECTION.NEW_PAGE)
    section_back.header.is_linked_to_previous = False
    section_back.footer.is_linked_to_previous = False
    
    # Zero margins for full bleed background
    section_back.top_margin = Inches(0)
    section_back.bottom_margin = Inches(0)
    section_back.left_margin = Inches(0)
    section_back.right_margin = Inches(0)
    section_back.page_width = Inches(8.5)
    section_back.page_height = Inches(11.0)
    
    back_table = doc.add_table(1, 1)
    back_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    back_table.autofit = False
    
    row_b = back_table.rows[0]
    row_b.height = Inches(11.0)
    row_b.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    cell_b = back_table.cell(0, 0)
    cell_b.width = Inches(8.5)
    set_cell_background(cell_b, HEX_MIDNIGHT_NAVY)
    set_cell_margins(cell_b, top=720, bottom=500, left=720, right=720)
    set_cell_valign(cell_b, "center")
    remove_table_borders(back_table)
    
    # Inside Back Cover Content
    p_b_logo = cell_b.paragraphs[0]
    set_para_format(p_b_logo, before_pt=0, after_pt=18, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if os.path.exists(logo_path):
        p_b_logo.add_run().add_picture(logo_path, width=Inches(3.2))
    else:
        r_b_logo_fail = p_b_logo.add_run("[ SMART IMPORT TRADE ]")
        set_font(r_b_logo_fail, "Montserrat SemiBold", 18, True, COLOR_GOLD_LIGHT)
        
    p_b_serv = cell_b.add_paragraph()
    set_para_format(p_b_serv, before_pt=0, after_pt=36, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    r_b_serv = p_b_serv.add_run("Importaciones  ·  Sourcing Internacional  ·  Desarrollo de Proveedores  ·  Comercio Exterior")
    set_font(r_b_serv, "Montserrat Medium", 9.5, False, COLOR_GOLD_LIGHT)
    
    # Large Mission Quote Box (using table in back cover)
    quote_table = cell_b.add_table(1, 1)
    quote_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(quote_table)
    
    cell_q = quote_table.cell(0, 0)
    cell_q.width = Inches(6.0)
    set_cell_background(cell_q, HEX_MIDNIGHT_NAVY)
    set_cell_margins(cell_q, top=100, bottom=100, left=100, right=100)
    set_cell_borders(cell_q, top="single", bottom="single", color=HEX_COPPER, size="12")
    
    p_q = cell_q.paragraphs[0]
    set_para_format(p_q, before_pt=12, after_pt=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    r_q = p_q.add_run(
        "\"Nuestra misión es identificar oportunidades de importación rentables, "
        "reducir riesgos comerciales y optimizar la inversión de nuestros clientes.\""
    )
    set_font(r_q, "Montserrat Medium", 13, False, COLOR_WARM_CREAM, italic=True)
    
    p_b_foot = cell_b.add_paragraph()
    set_para_format(p_b_foot, before_pt=48, after_pt=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    r_b_foot = p_b_foot.add_run("Documento confidencial – Uso exclusivo de Importaciones Manorca S.A.C.\n")
    set_font(r_b_foot, "Calibri", 8.5, False, COLOR_STEEL_GRAY)
    r_b_web = p_b_foot.add_run("www.smartimporttrade.com")
    set_font(r_b_web, "Montserrat Medium", 9, True, COLOR_COPPER)
    
    # Save Document
    out_path = "Informe_Preliminar_Importaciones_Manorca.docx"
    print(f"Saving document to {out_path}...")
    doc.save(out_path)
    print("Document saved successfully!")

if __name__ == "__main__":
    generate_docx()
